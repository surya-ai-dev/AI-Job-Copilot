"""Resume Parser Agent implementing Strategy Pattern for text extraction from PDF and DOCX."""

import logging
import re
import unicodedata
from abc import ABC, abstractmethod
from typing import BinaryIO, Dict, List
import xml.etree.ElementTree as ET

import pdfplumber
from docx import Document

from backend.app.ai.exceptions import (
    UnsupportedFileTypeError,
    TextExtractionError
)
from backend.app.ai.schemas.resume_parser_schema import ResumeParserResult

logger = logging.getLogger(__name__)

class BaseFormatParser(ABC):
    """Abstract base class for format-specific text extraction strategies."""

    @abstractmethod
    def extract_text_and_pages(self, file_stream: BinaryIO) -> tuple[str, int]:
        """Extracts text and page count from the given binary file stream.

        Args:
            file_stream (BinaryIO): Binary stream of the file content.

        Returns:
            tuple[str, int]: A tuple containing the raw extracted text and the page count.

        Raises:
            TextExtractionError: If text extraction fails.
        """
        pass


class PDFFormatParser(BaseFormatParser):
    """Format parser for PDF documents using pdfplumber."""

    def extract_text_and_pages(self, file_stream: BinaryIO) -> tuple[str, int]:
        try:
            # Ensure we are at the start of the stream
            file_stream.seek(0)
            with pdfplumber.open(file_stream) as pdf:
                pages = pdf.pages
                page_count = len(pages)
                
                if page_count == 0:
                    raise TextExtractionError("The PDF document has 0 pages.")

                extracted_pages = []
                for i, page in enumerate(pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_pages.append(page_text)
                    else:
                        logger.warning("No text extracted from page %d of the PDF.", i + 1)
                
                raw_text = "\n".join(extracted_pages)
                return raw_text, page_count
        except Exception as e:
            logger.error("Failed to parse PDF document: %s", str(e), exc_info=True)
            raise TextExtractionError(f"PDF parsing error: {str(e)}") from e


class DocxFormatParser(BaseFormatParser):
    """Format parser for DOCX documents using python-docx."""

    def extract_text_and_pages(self, file_stream: BinaryIO) -> tuple[str, int]:
        try:
            # Ensure we are at the start of the stream
            file_stream.seek(0)
            doc = Document(file_stream)
            
            text_parts = []
            
            # Extract paragraph texts
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Extract table cell texts
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Avoid duplicating cell texts in case of merged cells
                        cell_text = cell.text.strip()
                        if cell_text and (not row_text or row_text[-1] != cell_text):
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            raw_text = "\n".join(text_parts)
            
            # Extract page count from metadata (docProps/app.xml) if available, fallback to 1
            page_count = self._get_page_count(doc)
            
            return raw_text, page_count
        except Exception as e:
            logger.error("Failed to parse DOCX document: %s", str(e), exc_info=True)
            raise TextExtractionError(f"DOCX parsing error: {str(e)}") from e

    def _get_page_count(self, doc: Document) -> int:
        """Attempts to read the page count from DOCX package app.xml properties."""
        try:
            for part in doc.part.package.parts:
                if part.partname == "/docProps/app.xml":
                    root = ET.fromstring(part.blob)
                    # OpenXML schema namespace
                    ns = {"extended": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}
                    pages_elem = root.find(".//extended:Pages", ns)
                    if pages_elem is not None and pages_elem.text:
                        return int(pages_elem.text)
        except Exception as e:
            logger.warning("Could not extract page count from DOCX properties: %s. Defaulting to 1.", str(e))
        return 1


class ResumeParserAgent:
    """Agent responsible for orchestrating the parsing of resume files.

    Uses dependency injection to manage format parsers and applies text cleaning rules.
    """

    def __init__(self, format_parsers: Dict[str, BaseFormatParser] = None):
        """Initializes the agent with supported format parsers.

        Args:
            format_parsers (Dict[str, BaseFormatParser], optional): Mapping of file extensions
                to their respective format parsers. Defaults to supporting PDF and DOCX.
        """
        if format_parsers is None:
            self._parsers = {
                "pdf": PDFFormatParser(),
                "docx": DocxFormatParser()
            }
        else:
            self._parsers = format_parsers

    def clean_text(self, text: str) -> str:
        """Cleans and normalizes extracted text.

        Removes redundant spacing, normalizes unicode characters, and standardizes newlines.
        """
        if not text:
            return ""
        
        # Normalize unicode sequences
        text = unicodedata.normalize("NFKC", text)
        
        # Standardize line endings to \n
        text = re.sub(r"\r\n?", "\n", text)
        
        # Replace tabs with spaces
        text = text.replace("\t", " ")
        
        cleaned_lines = []
        for line in text.split("\n"):
            # Clean consecutive whitespaces within the line
            cleaned_line = re.sub(r"\s+", " ", line).strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
                
        return "\n".join(cleaned_lines)

    def parse(self, file_stream: BinaryIO, file_type: str) -> ResumeParserResult:
        """Parses the given resume binary stream according to its file type.

        Args:
            file_stream (BinaryIO): Binary stream of the file content.
            file_type (str): The file extension/type (e.g. 'pdf', 'docx').

        Returns:
            ResumeParserResult: Structured object containing parsed results.

        Raises:
            UnsupportedFileTypeError: If the file type is not supported.
            TextExtractionError: If text extraction fails or returns empty text.
        """
        normalized_type = file_type.lower().strip().replace(".", "")
        parser = self._parsers.get(normalized_type)
        
        if not parser:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: '{file_type}'. Supported types: {list(self._parsers.keys())}"
            )

        logger.info("Parsing resume with file type: %s", normalized_type)
        raw_text, page_count = parser.extract_text_and_pages(file_stream)
        
        cleaned_text = self.clean_text(raw_text)
        
        if not cleaned_text.strip():
            raise TextExtractionError(
                "Text extraction succeeded but no readable text content was resolved."
            )

        character_count = len(cleaned_text)
        warnings = []

        # Heuristic checks for warnings
        if character_count < 100:
            warnings.append(
                "Extracted text is extremely short (< 100 chars). The file might be an image, scanned PDF, or corrupted."
            )
        if "\uFFFD" in raw_text:
            warnings.append(
                "Extracted text contains unicode replacement characters (\\uFFFD), indicating possible decoding issues."
            )

        return ResumeParserResult(
            raw_text=cleaned_text,
            page_count=page_count,
            file_type=normalized_type,
            character_count=character_count,
            warnings=warnings
            # No AI extraction, skill extraction, or experience chunking here, per rules
        )
